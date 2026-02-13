import sys
import re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

# Read the docx file
doc = Document(r'C:\Users\fujitsu\Desktop\fly-book\docs\ПВП ДАУ наказ №2 від 05.01.2015.docx')

def clean_text(text):
    text = re.sub(r'\{[^}]*\}', '', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def extract_content(start_line, end_line):
    content = []
    for i in range(start_line, min(end_line, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if text:
            content.append(text)
    return '\n\n'.join(content)

# Generate SQL for Section IV, Chapter 4
title4 = '4. Підготовка до польотів авіаційного персоналу'
parent_id4 = 755
order_num4 = 4
start4 = 510
end4 = 589

content4 = extract_content(start4, end4)
cleaned4 = clean_text(content4)

sql4 = f"""INSERT INTO guide_sections (document_id, title, content, parent_id, order_num)
VALUES (1, '{title4}', $${cleaned4}$$, {parent_id4}, {order_num4});"""

# Write to file
with open('section_iv_chapter4.sql', 'w', encoding='utf-8') as f:
    f.write(sql4)

print(f"Generated SQL for Chapter 4 (length: {len(cleaned4)} chars)")
print(f"Saved to section_iv_chapter4.sql")

# Generate SQL for Section IV, Chapter 5
title5 = '5. Розвідка погоди'
parent_id5 = 755
order_num5 = 5
start5 = 590
end5 = 655

content5 = extract_content(start5, end5)
cleaned5 = clean_text(content5)

sql5 = f"""INSERT INTO guide_sections (document_id, title, content, parent_id, order_num)
VALUES (1, '{title5}', $${cleaned5}$$, {parent_id5}, {order_num5});"""

# Write to file
with open('section_iv_chapter5.sql', 'w', encoding='utf-8') as f:
    f.write(sql5)

print(f"Generated SQL for Chapter 5 (length: {len(cleaned5)} chars)")
print(f"Saved to section_iv_chapter5.sql")

print("\nContent preview Chapter 4 (first 500 chars):")
print(cleaned4[:500])
print("\nContent preview Chapter 5 (first 500 chars):")
print(cleaned5[:500])

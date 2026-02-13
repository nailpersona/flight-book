import requests
import json
import sys
import re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'C:\Users\fujitsu\Desktop\fly-book\docs\ПВП ДАУ наказ №2 від 05.01.2015.docx')

SUPABASE_URL = 'https://klqxadvtvxvizgdjmegx.supabase.co'
SUPABASE_KEY = 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImZseWJvb2siLCJleHAiOjE3NDk3MTg4MDAsImVhbXAiOjE3NDk3MjQ4MDAsIn0.zY5H_Mo5SxqPo7IMdCOg-nH0d1HIBNPEYQ-JAjSDE'

def clean_text(text):
    """Remove legal notes in curly braces and extra whitespace"""
    text = re.sub(r'\{[^}]*\}', '', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def extract_content(start_line, end_line):
    content = []
    for i in range(start_line, min(end_line, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if text:
            content.append(text)
    return '\n\n'.join(content)

# Section III: Група керівництва польотами (id: 754)
section_iii_id = 754

section_iii_chapters = [
    ('2. Допуск осіб групи керівництва польотами до керівництва польотами', 2, 408, 455),
    ('3. Перевірки осіб групи керівництва польотами', 3, 446, 456),
    ('4. Допустимі перерви і порядок відновлення навичок осіб групи керівництва польотами в керівництві польотами', 4, 459, 463)
]

# Section IV: Організація польотів (id: 755)
section_iv_id = 755

section_iv_chapters = [
    ('1. Зміст організації польотів', 1, 466, 471),
    ('2. Прийняття рішення та постановка завдань на польоти', 2, 472, 495),
    ('3. Планування польотів', 3, 496, 509),
    ('4. Підготовка до польотів авіаційного персоналу', 4, 510, 589),
    ('5. Розвідка погоди', 5, 590, 655)
]

all_chapters = []
for title, order, start, end in section_iii_chapters:
    content = extract_content(start, end)
    cleaned = clean_text(content)
    all_chapters.append({
        'document_id': 1,
        'title': title,
        'content': cleaned,
        'parent_id': section_iii_id,
        'order_num': order
    })

for title, order, start, end in section_iv_chapters:
    content = extract_content(start, end)
    cleaned = clean_text(content)
    all_chapters.append({
        'document_id': 1,
        'title': title,
        'content': cleaned,
        'parent_id': section_iv_id,
        'order_num': order
    })

# Insert via Supabase REST API
headers = {
    'apikey': SUPABASE_KEY,
    'Authorization': f'Bearer {SUPABASE_KEY}',
    'Content-Type': 'application/json'
}

for chapter in all_chapters:
    response = requests.post(
        f'{SUPABASE_URL}/rest/v1/guide_sections',
        headers=headers,
        json=chapter
    )
    if response.status_code == 201:
        print(f"Inserted: {chapter['title']}")
    else:
        print(f"Error inserting {chapter['title']}: {response.status_code}")
        print(response.text)

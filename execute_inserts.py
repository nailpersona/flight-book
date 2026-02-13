import psycopg2
import sys
import re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

# Database connection
conn = psycopg2.connect(
    host='db.klqxadvtvxvizgdjmegx.supabase.co',
    database='postgres',
    user='postgres',
    password='(your_password)',  # Will be replaced
    port='5432'
)

# Read the docx file
doc = Document(r'C:\Users\fujitsu\Desktop\fly-book\docs\ПВП ДАУ наказ №2 від 05.01.2015.docx')

def clean_text(text):
    text = re.sub(r'\{[^}]*\}', '', text)
    text = re.sub(r('\n+', '\n', text)
    return text.strip()

def extract_content(start_line, end_line):
    content = []
    for i in range(start_line, min(end_line, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if text:
            content.append(text)
    return '\n\n'.join(content)

# Data for remaining chapters
sections_data = [
    # Section III chapters (parent_id = 754)
    ('2. Допуск осіб групи керівництва польотами до керівництва польотами', 754, 2, 408, 455),
    ('3. Перевірки осіб групи керівництва польотами', 754, 3, 446, 456),
    ('4. Допустимі перерви і порядок відновлення навичок осіб групи керівництва польотами в керівництві польотами', 754, 4, 459, 463),
    # Section IV chapters (parent_id = 755)
    ('2. Прийняття рішення та постановка завдань на польоти', 755, 2, 472, 495),
    ('3. Планування польотів', 755, 3, 496, 509),
    ('4. Підготовка до польотів авіаційного персоналу', 755, 4, 510, 589),
    ('5. Розвідка погоди', 755, 5, 590, 655)
]

# Generate SQL for each section
for title, parent_id, order_num, start, end in sections_data:
    content = extract_content(start, end)
    cleaned = clean_text(content)

    sql = "INSERT INTO guide_sections (document_id, title, content, parent_id, order_num) VALUES (1, %s, %s, %s, %s);"

    print(f"Would insert: {title}")
    # Note: You'll need to execute this via Supabase SQL Editor instead
    # The connection requires the DB password which we don't have

print("\nInstead, execute the following SQL in Supabase SQL Editor:")
print("(Each section is already generated in separate files)")

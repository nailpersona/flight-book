import sys
import re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'C:\Users\fujitsu\Desktop\fly-book\docs\ПВП ДАУ наказ №2 від 05.01.2015.docx')

def clean_text(text):
    """Remove legal notes in curly braces and extra whitespace"""
    text = re.sub(r'\{[^}]*\}', '', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def extract_content(start_line, end_line):
    content = []
    for i in range(start_line, min(end_line, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if text:
            content.append(text)
    return '\n\n'.join(content)

# Section III: Група керівництва польотами (id: 754)
section_iii_id = 754

section_iii_chapters = [
    ('1. Склад групи керівництва польотами', 1, 317, 407),
    ('2. Допуск осіб групи керівництва польотами до керівництва польотами', 2, 408, 455),
    ('3. Перевірки осіб групи керівництва польотами', 3, 446, 456),
    ('4. Допустимі перерви і порядок відновлення навичок осіб групи керівництва польотами в керівництві польотами', 4, 459, 463)
]

# Section IV: Організація польотів (id: 755)
section_iv_id = 755

section_iv_chapters = [
    ('1. Зміст організації польотів', 1, 466, 471),
    ('2. Прийняття рішення та постановка завдань на польоти', 2, 472, 495),
    ('3. Планування польотів', 3, 496, 509),
    ('4. Підготовка до польотів авіаційного персоналу', 4, 510, 589),
    ('5. Розвідка погоди', 5, 590, 655)
]

# Generate SQL for Section III
print("-- ============================================")
print("-- РОЗДІЛ III. Група керівництва польотами")
print("-- ============================================")
print()

for title, order, start, end in section_iii_chapters:
    content = extract_content(start, end)
    cleaned = clean_text(content)
    print(f"-- {title}")
    print(f"INSERT INTO guide_sections (document_id, title, content, parent_id, order_num)")
    print(f"VALUES (1, '{title}', $$")
    print(cleaned)
    print(f"$$, {section_iii_id}, {order});")
    print()

print()
print("-- ============================================")
print("-- РОЗДІЛ IV. Організація польотів")
print("-- ============================================")
print()

for title, order, start, end in section_iv_chapters:
    content = extract_content(start, end)
    cleaned = clean_text(content)
    print(f"-- {title}")
    print(f"INSERT INTO guide_sections (document_id, title, content, parent_id, order_num)")
    print(f"VALUES (1, '{title}', $$")
    print(cleaned)
    print(f"$$, {section_iv_id}, {order});")
    print()

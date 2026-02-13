import sys
import json
import re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

# Read the docx file
doc = Document(r'C:\Users\fujitsu\Desktop\fly-book\docs\ПВП ДАУ наказ №2 від 05.01.2015.docx')

def clean_text(text):
    text = re.sub(r'\{[^}]*\}', '', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def extract_content(start_line, end_line):
    content = []
    for i in range(start_line, min(end_line, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if text:
            content.append(text)
    return '\n\n'.join(content)

# Define all sections to insert
sections_data = [
    # Section III chapters (parent_id = 754)
    {
        'title': '2. Допуск осіб групи керівництва польотами до керівництва польотами',
        'parent_id': 754,
        'order_num': 2,
        'start': 408,
        'end': 455
    },
    {
        'title': '3. Перевірки осіб групи керівництва польотами',
        'parent_id': 754,
        'order_num': 3,
        'start': 446,
        'end': 456
    },
    {
        'title': '4. Допустимі перерви і порядок відновлення навичок осіб групи керівництва польотами в керівництві польотами',
        'parent_id': 754,
        'order_num': 4,
        'start': 459,
        'end': 463
    },
    # Section IV chapters (parent_id = 755)
    {
        'title': '2. Прийняття рішення та постановка завдань на польоти',
        'parent_id': 755,
        'order_num': 2,
        'start': 472,
        'end': 495
    },
    {
        'title': '3. Планування польотів',
        'parent_id': 755,
        'order_num': 3,
        'start': 496,
        'end': 509
    },
    {
        'title': '4. Підготовка до польотів авіаційного персоналу',
        'parent_id': 755,
        'order_num': 4,
        'start': 510,
        'end': 589
    },
    {
        'title': '5. Розвідка погоди',
        'parent_id': 755,
        'order_num': 5,
        'start': 590,
        'end': 655
    }
]

# Build the SQL statements
sql_statements = []

for section in sections_data:
    content = extract_content(section['start'], section['end'])
    cleaned = clean_text(content)
    parent_id = section['parent_id']
    order_num = section['order_num']
    title = section['title']

    sql = f"""INSERT INTO guide_sections (document_id, title, content, parent_id, order_num)
VALUES (1, '{title}', $${cleaned}$, {parent_id}, {order_num});"""

    sql_statements.append(sql)

# Write all SQL to a single file
with open('insert_all_sections.sql', 'w', encoding='utf-8') as f:
    f.write('\n\n'.join(sql_statements))

print(f"Generated {len(sql_statements)} SQL statements in insert_all_sections.sql")
print("Copy the content and execute in Supabase SQL Editor.")

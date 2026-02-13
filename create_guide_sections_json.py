import sys
import re
import json
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'C:\Users\fujitsu\Desktop\fly-book\docs\ПВП ДАУ наказ №2 від 05.01.2015.docx')

def clean_text(text):
    """Remove legal notes in curly braces and extra whitespace"""
    text = re.sub(r'\{[^}]*\}', '', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def extract_content(start_line, end_line):
    content = []
    for i in range(start_line, min(end_line, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if text:
            content.append(text)
    return '\n\n'.join(content)

# Section III: Група керівництва польотами
section_iii_chapters = [
    {
        'title': '1. Склад групи керівництва польотами',
        'order': 1,
        'start': 317,
        'end': 407
    },
    {
        'title': '2. Допуск осіб групи керівництва польотами до керівництва польотами',
        'order': 2,
        'start': 408,
        'end': 455
    },
    {
        'title': '3. Перевірки осіб групи керівництва польотами',
        'order': 3,
        'start': 446,
        'end': 456
    },
    {
        'title': '4. Допустимі перерви і порядок відновлення навичок осіб групи керівництва польотами в керівництві польотами',
        'order': 4,
        'start': 459,
        'end': 463
    }
]

# Section IV: Організація польотів
section_iv_chapters = [
    {
        'title': '1. Зміст організації польотів',
        'order': 1,
        'start': 466,
        'end': 471
    },
    {
        'title': '2. Прийняття рішення та постановка завдань на польоти',
        'order': 2,
        'start': 472,
        'end': 495
    },
    {
        'title': '3. Планування польотів',
        'order': 3,
        'start': 496,
        'end': 509
    },
    {
        'title': '4. Підготовка до польотів авіаційного персоналу',
        'order': 4,
        'start': 510,
        'end': 589
    },
    {
        'title': '5. Розвідка погоди',
        'order': 5,
        'start': 590,
        'end': 655
    }
]

# Build sections data
sections = []

# Section III
section_iii = {
    'title': 'III. Група керівництва польотами',
    'order_num': 3,
    'chapters': []
}

for ch in section_iii_chapters:
    content = extract_content(ch['start'], ch['end'])
    section_iii['chapters'].append({
        'title': ch['title'],
        'content': clean_text(content),
        'order_num': ch['order']
    })

sections.append(section_iii)

# Section IV
section_iv = {
    'title': 'IV. Організація польотів',
    'order_num': 4,
    'chapters': []
}

for ch in section_iv_chapters:
    content = extract_content(ch['start'], ch['end'])
    section_iv['chapters'].append({
        'title': ch['title'],
        'content': clean_text(content),
        'order_num': ch['order']
    })

sections.append(section_iv)

# Output JSON
print(json.dumps({'sections': sections}, ensure_ascii=False, indent=2))

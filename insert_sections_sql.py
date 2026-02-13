import sys
import re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'C:\Users\fujitsu\Desktop\fly-book\docs\ПВП ДАУ наказ №2 від 05.01.2015.docx')

def clean_text(text):
    """Remove legal notes in curly braces and extra whitespace"""
    text = re.sub(r'\{[^}]*\}', '', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def extract_content(start_line, end_line):
    content = []
    for i in range(start_line, min(end_line, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if text:
            content.append(text)
    return '\n\n'.join(content)

def escape_sql_string(text):
    """Escape single quotes for SQL"""
    return text.replace("'", "''")

# Section III: Група керівництва польотами (id: 754)
section_iii_id = 754
section_iii_chapters = [
    {
        'title': '1. Склад групи керівництва польотами',
        'order': 1,
        'start': 317,
        'end': 407
    },
    {
        'title': '2. Допуск осіб групи керівництва польотами до керівництва польотами',
        'order': 2,
        'start': 408,
        'end': 455
    },
    {
        'title': '3. Перевірки осіб групи керівництва польотами',
        'order': 3,
        'start': 446,
        'end': 456
    },
    {
        'title': '4. Допустимі перерви і порядок відновлення навичок осіб групи керівництва польотами в керівництві польотами',
        'order': 4,
        'start': 459,
        'end': 463
    }
]

# Section IV: Організація польотів (id: 755)
section_iv_id = 755
section_iv_chapters = [
    {
        'title': '1. Зміст організації польотів',
        'order': 1,
        'start': 466,
        'end': 471
    },
    {
        'title': '2. Прийняття рішення та постановка завдань на польоти',
        'order': 2,
        'start': 472,
        'end': 495
    },
    {
        'title': '3. Планування польотів',
        'order': 3,
        'start': 496,
        'end': 509
    },
    {
        'title': '4. Підготовка до польотів авіаційного персоналу',
        'order': 4,
        'start': 510,
        'end': 589
    },
    {
        'title': '5. Розвідка погоди',
        'order': 5,
        'start': 590,
        'end': 655
    }
]

# Generate SQL for Section III
print("-- ============================================")
print("-- РОЗДІЛ III. Група керівництва польотами (parent_id = 754)")
print("-- ============================================")
print()

for ch in section_iii_chapters:
    content = extract_content(ch['start'], ch['end'])
    cleaned = clean_text(content)
    # Write to file for each chapter
    filename = f"section_iii_chapter_{ch['order']}.sql"
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(f"-- {ch['title']}\n")
        f.write(f"INSERT INTO guide_sections (document_id, title, content, parent_id, order_num)\n")
        f.write(f"VALUES (1, '{ch['title']}', $$\n")
        f.write(cleaned)
        f.write(f"\n$$, {section_iii_id}, {ch['order']});\n")
    print(f"-- Generated: {filename}")

print()
print("-- ============================================")
print("-- РОЗДІЛ IV. Організація польотів (parent_id = 755)")
print("-- ============================================")
print()

for ch in section_iv_chapters:
    content = extract_content(ch['start'], ch['end'])
    cleaned = clean_text(content)
    # Write to file for each chapter
    filename = f"section_iv_chapter_{ch['order']}.sql"
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(f"-- {ch['title']}\n")
        f.write(f"INSERT INTO guide_sections (document_id, title, content, parent_id, order_num)\n")
        f.write(f"VALUES (1, '{ch['title']}', $$\n")
        f.write(cleaned)
        f.write(f"\n$$, {section_iv_id}, {ch['order']});\n")
    print(f"-- Generated: {filename}")

print()
print("-- All files generated. Execute them in Supabase SQL Editor.")

import sys
import re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'C:\Users\fujitsu\Desktop\fly-book\docs\ПВП ДАУ наказ №2 від 05.01.2015.docx')

# Function to clean text - remove legal notes in curly braces
def clean_text(text):
    # Remove content in curly braces like {Наказ...}
    text = re.sub(r'\{[^}]*\}', '', text)
    # Remove extra whitespace
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

# Function to extract content between two line numbers
def extract_content(start_line, end_line):
    content = []
    for i in range(start_line, min(end_line, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if text:
            content.append(text)
    return '\n\n'.join(content)

# Section III: Група керівництва польотами
section_iii_chapters = {
    '1. Склад групи керівництва польотами': (317, 407),
    '2. Допуск осіб групи керівництва польотами до керівництва польотами': (408, 455),
    '3. Перевірки осіб групи керівництва польотами': (446, 456),
    '4. Допустимі перерви і порядок відновлення навичок осіб групи керівництва польотами в керівництві польотами': (459, 463)
}

# Section IV: Організація польотів
section_iv_chapters = {
    '1. Зміст організації польотів': (466, 471),
    '2. Прийняття рішення та постановка завдань на польоти': (472, 495),
    '3. Планування польотів': (496, 509),
    '4. Підготовка до польотів авіаційного персоналу': (510, 589),
    '5. Розвідка погоди': (590, 655)
}

# Generate SQL for Section III
print("-- ============================================")
print("-- РОЗДІЛ III. Група керівництва польотами")
print("-- ============================================")
print()

# First, insert the main section (accordion - no content)
print("-- Step 1: Insert main section III (accordion)")
print("INSERT INTO guide_sections (document_id, title, parent_id, order_num)")
print("VALUES (1, 'III. Група керівництва польотами', NULL, 3);")
print()

# Get the parent_id for section III (we'll need to query this after insertion)
print("-- Note: After insertion, get the ID of this section to use as parent_id for subsections")
print("-- SELECT id FROM guide_sections WHERE title = 'III. Група керівництва польотами';")
print("-- Assuming parent_id = XXX for Section III")
print()

for i, (chapter_title, (start, end)) in enumerate(section_iii_chapters.items(), 1):
    print(f"-- Chapter {i}: {chapter_title}")
    content = extract_content(start, end)
    cleaned_content = clean_text(content)
    # Escape single quotes for SQL
    escaped_content = cleaned_content.replace("'", "''")

    print("INSERT INTO guide_sections (document_id, title, content, parent_id, order_num)")
    print(f"VALUES (1, '{chapter_title}', $$")
    print(escaped_content)
    print("$$, XXX, " + str(i) + ");")
    print()

print()
print("-- ============================================")
print("-- РОЗДІЛ IV. Організація польотів")
print("-- ============================================")
print()

# First, insert the main section IV (accordion - no content)
print("-- Step 1: Insert main section IV (accordion)")
print("INSERT INTO guide_sections (document_id, title, parent_id, order_num)")
print("VALUES (1, 'IV. Організація польотів', NULL, 4);")
print()

print("-- Note: After insertion, get the ID of this section to use as parent_id for subsections")
print("-- SELECT id FROM guide_sections WHERE title = 'IV. Організація польотів';")
print("-- Assuming parent_id = YYY for Section IV")
print()

for i, (chapter_title, (start, end)) in enumerate(section_iv_chapters.items(), 1):
    print(f"-- Chapter {i}: {chapter_title}")
    content = extract_content(start, end)
    cleaned_content = clean_text(content)
    # Escape single quotes for SQL
    escaped_content = cleaned_content.replace("'", "''")

    print("INSERT INTO guide_sections (document_id, title, content, parent_id, order_num)")
    print(f"VALUES (1, '{chapter_title}', $$")
    print(escaped_content)
    print("$$, YYY, " + str(i) + ");")
    print()

import sys
import json
import re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'C:\Users\fujitsu\Desktop\fly-book\docs\ПВП ДАУ наказ №2 від 05.01.2015.docx')

def clean_text(text):
    text = re.sub(r'\{[^}]*\}', '', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def extract_content(start_line, end_line):
    content = []
    for i in range(start_line, min(end_line, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if text:
            content.append(text)
    return '\n\n'.join(content)

# Data for remaining chapters
sections = [
    # Section III chapters (parent_id = 754)
    {
        'title': '2. Допуск осіб групи керівництва польотами до керівництва польотами',
        'parent_id': 754,
        'order_num': 2,
        'start': 408,
        'end': 455
    },
    {
        'title': '3. Перевірки осіб групи керівництва польотами',
        'parent_id': 754,
        'order_num': 3,
        'start': 446,
        'end': 456
    },
    {
        'title': '4. Допустимі перерви і порядок відновлення навичок осіб групи керівництва польотами в керівництві польотами',
        'parent_id': 754,
        'order_num': 4,
        'start': 459,
        'end': 463
    },
    # Section IV chapters (parent_id = 755)
    {
        'title': '2. Прийняття рішення та постановка завдань на польоти',
        'parent_id': 755,
        'order_num': 2,
        'start': 472,
        'end': 495
    },
    {
        'title': '3. Планування польотів',
        'parent_id': 755,
        'order_num': 3,
        'start': 496,
        'end': 509
    },
    {
        'title': '4. Підготовка до польотів авіаційного персоналу',
        'parent_id': 755,
        'order_num': 4,
        'start': 510,
        'end': 589
    },
    {
        'title': '5. Розвідка погоди',
        'parent_id': 755,
        'order_num': 5,
        'start': 590,
        'end': 655
    }
]

# Generate individual SQL files for each section
for section in sections:
    content = extract_content(section['start'], section['end'])
    cleaned = clean_text(content)

    # Write to SQL file
    filename = f"insert_{section['parent_id']}_chapter_{section['order_num']}.sql"
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(f"-- {section['title']}\n")
        f.write(f"INSERT INTO guide_sections (document_id, title, content, parent_id, order_num)\n")
        f.write(f"VALUES (1, '{section['title']}', $$\n")
        f.write(cleaned)
        f.write(f"\n$$, {section['parent_id']}, {section['order_num']});\n")

    print(f"Generated: {filename}")

print()
print("All SQL files generated successfully!")
print("Execute each file in Supabase SQL Editor.")
